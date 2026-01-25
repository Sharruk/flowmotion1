from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    doc = Document()

    # --- TITLE PAGE ---
    title = doc.add_heading('Sri Sivasubramaniya Nadar College of Engineering', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(An Autonomous Institution, Affiliated to Anna University, Chennai)\nKalavakkam - 603 110')
    run.font.size = Pt(12)

    doc.add_paragraph('\n\n\n')
    
    proj_title = doc.add_paragraph()
    proj_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = proj_title.add_run('LAB RECORD DOCUMENTATION\nFlowMotion: Intelligent Habit & Routine Management')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph('\n\n')

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = details.add_run('Course: ICS1405 – Software Engineering Principles and Practices\n'
                         'Regulation: 2023\n'
                         'Batch: 2024–2029\n'
                         'Academic Year: 2026–2027 (Even Semester)')
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- EXPERIMENT 1 ---
    doc.add_heading('EXPERIMENT 1: Problem Identification, Requirements Analysis, System Design & Implementation', level=1)
    
    doc.add_heading('1. Problem Identification', level=2)
    doc.add_paragraph('Problem Statement: Modern users, especially students and professionals on desktop environments, struggle with maintaining consistent habits due to digital distractions and lack of integrated reminder systems.')
    doc.add_paragraph('Real-world Motivation: Habit tracking is often confined to mobile devices, which are secondary to the primary workstation (Linux Desktop). FlowMotion aims to bridge this gap by integrating habit tracking directly into the desktop workflow.')
    doc.add_paragraph('Limitations of Existing Systems: Most habit trackers are mobile-only, rely on browser-based push notifications (easily missed), or lack genuine emotional engagement, leading to "notification fatigue".')

    doc.add_heading('2. Stakeholder Analysis', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stakeholder'
    hdr_cells[1].text = 'Role'
    hdr_cells[2].text = 'Motive'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'End Users'
    row_cells[1].text = 'Primary Consumer'
    row_cells[2].text = 'Consistency in habits & productivity.'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Developers'
    row_cells[1].text = 'System Architect'
    row_cells[2].text = 'Implementation of SE principles.'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Academic Evaluators'
    row_cells[1].text = 'Quality Assurance'
    row_cells[2].text = 'Assessment of modularity and logic.'

    doc.add_heading('3. Requirements Analysis', level=2)
    doc.add_heading('Functional Requirements', level=3)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'ID'
    table.rows[0].cells[1].text = 'Requirement'
    reqs = [('FR1', 'User authentication and profile management.'),
            ('FR2', 'Creation of Yes/No and Measurable habits.'),
            ('FR3', 'Automated background notifications via APScheduler.'),
            ('FR4', 'AI-powered emotional feedback based on performance.')]
    for rid, req in reqs:
        row = table.add_row().cells
        row[0].text = rid
        row[1].text = req

    doc.add_heading('4. System Design', level=2)
    doc.add_paragraph('Architecture: The system follows a Model-View-Template (MVT) pattern using Django.')
    doc.add_paragraph('Modules:\n- Frontend: HTML/JS for dashboard and habit tracking.\n- Backend: Django server managing logic and scheduling.\n- Database: SQLite3 for persistent storage.\n- AI Engine: Gemini AI integration for emotional feedback.')

    doc.add_heading('5. Implementation', level=2)
    doc.add_paragraph('The prototype is built using Python 3.11 and Django 5.x. Background tasks are handled by APScheduler, while notifications are routed through the Linux "notify-send" utility for native integration.')

    doc.add_page_break()

    # --- EXPERIMENT 2 ---
    doc.add_heading('EXPERIMENT 2: Software Requirements Specification (SRS)', level=1)
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph('Purpose: This SRS defines the functional and non-functional requirements for FlowMotion v1.0.')
    doc.add_paragraph('Scope: Focuses on local Linux desktop users with Ubuntu/NixOS.')

    doc.add_heading('2. Functional Requirements', level=2)
    doc.add_paragraph('1. The system shall allow users to register and login securely.')
    doc.add_paragraph('2. The system shall send a pre-reminder 10 minutes before a scheduled habit.')
    doc.add_paragraph('3. The system shall repeat missed reminders every 10 minutes until completion.')
    doc.add_paragraph('4. The system shall generate unique AI feedback messages for each habit completion.')

    doc.add_heading('3. Non-Functional Requirements', level=2)
    doc.add_paragraph('- Performance: Notifications must be triggered within 60 seconds of the scheduled time.')
    doc.add_paragraph('- Usability: Habit completion should require no more than two clicks from the dashboard.')
    doc.add_paragraph('- Reliability: The scheduler must restart automatically with the web server.')

    doc.add_heading('4. Constraints', level=2)
    doc.add_paragraph('- Platform: Must run on Linux with "notify-send" available.')
    doc.add_paragraph('- Database: SQLite3 for portability in academic environments.')

    doc.add_page_break()

    # --- EXPERIMENT 3 ---
    doc.add_heading('EXPERIMENT 3: Data Flow Diagrams (DFD)', level=1)
    doc.add_heading('1. Level-0 DFD (Context Diagram)', level=2)
    doc.add_paragraph('Description: The User interacts with the FlowMotion System by providing Habit Details and Completion Status. The system provides Notifications and Analytics.')
    
    doc.add_heading('2. Level-1 DFD', level=2)
    doc.add_paragraph('Processes:\n1.0 User Authentication\n2.0 Habit Management\n3.0 Background Scheduler\n4.0 AI Feedback Engine')
    doc.add_paragraph('Data Stores: User DB, Habit DB, Response DB.')

    doc.add_heading('3. Level-2 DFD (Notification Module)', level=2)
    doc.add_paragraph('Logic Flow:\n1. Scheduler checks Habit DB every 1 min.\n2. Logic compares Current Time vs Habit Time.\n3. If criteria met, send data to Notify-Send utility.')

    doc.add_heading('4. Conclusion', level=2)
    doc.add_paragraph('The DFDs successfully model the data movement from user input to background automation and AI feedback, ensuring a logical flow throughout the system.')

    doc.save('FlowMotion_Lab_Documentation.docx')
    print('Documentation generated: FlowMotion_Lab_Documentation.docx')

if __name__ == '__main__':
    create_documentation()
